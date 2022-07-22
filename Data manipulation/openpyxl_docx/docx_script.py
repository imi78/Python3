from docx import Document
from docx.shared import Pt

import openpyxl
document = Document()
# read excel sheet
wb = openpyxl.load_workbook('list.xlsx')
ws = wb['Protokol_list']

# iterate through the cells
for sheet in ws.iter_rows(min_col=3, max_col=3):
  for cell in sheet:
    if cell.value in [None,'Rank, name','Position']:
      continue
    else:
    
      p1 = document.add_paragraph()
      # Change the font and size of the text and the spacing
      p1.add_run('Dear ').font.name = 'Arial'
      p1.add_run().font.size = Pt(12)
      p1.paragraph_format.space_after = Pt(6)
      p1.paragraph_format.line_spacing = None
     
      p1.add_run(cell.value+ ',').font.name = 'Arial'
      
      p2 = document.add_paragraph()
      
      p2.add_run().font.size = Pt(12)
      p2.paragraph_format.space_after = Pt(6)
      p2.paragraph_format.line_spacing = None
      p2.add_run('You are cordially invited to NFIU BGR 7th anniversary' ).font.name = 'Arial'
      
      document.add_paragraph()
      document.add_paragraph()     


document.save('test.docx')